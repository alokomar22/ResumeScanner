from flask import Flask, request, render_template, jsonify
import os
import re
import PyPDF2
from docx import Document
import tempfile
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file with enhanced header detection"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                print(f"DEBUG: Page {page_num + 1} text length: {len(page_text)}")
                if page_text.strip():
                    # Split into lines and show first few lines (likely header)
                    lines = page_text.split('\n')
                    print(f"DEBUG: First 5 lines of page {page_num + 1}:")
                    for i, line in enumerate(lines[:5]):
                        if line.strip():
                            print(f"  Line {i+1}: {line.strip()}")
                    
                    text += page_text + "\n"
                
                # Try alternative extraction methods if standard method fails
                if not page_text.strip():
                    print(f"DEBUG: Trying alternative extraction for page {page_num + 1}")
                    try:
                        # Try to extract with different methods
                        if hasattr(page, 'extractText'):
                            alt_text = page.extractText()
                            if alt_text.strip():
                                text += alt_text + "\n"
                                print(f"DEBUG: Alternative extraction successful: {len(alt_text)} chars")
                    except:
                        pass
                        
    except Exception as e:
        print(f"Error reading PDF: {e}")
    
    print(f"DEBUG: Total PDF text extracted: {len(text)} characters")
    
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file including headers and footers"""
    text = ""
    try:
        doc = Document(file_path)
        
        # Extract text from headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    print(f"DEBUG: Header text: {paragraph.text}")
        
        # Extract text from main document
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    print(f"DEBUG: Footer text: {paragraph.text}")
        
        # Extract text from tables (if any) with enhanced structure handling
        for table_num, table in enumerate(doc.tables):
            print(f"DEBUG: Processing table {table_num + 1}")
            
            # Try to detect if this is a contact info table or structured resume table
            table_text = ""
            structured_data = {}
            
            for row_num, row in enumerate(table.rows):
                row_text = []
                for cell_num, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                        table_text += cell_text + " "
                        print(f"DEBUG: Table {table_num + 1}, Row {row_num + 1}, Cell {cell_num + 1}: '{cell_text}'")
                        
                        # Check if this cell contains contact information
                        cell_lower = cell_text.lower()
                        if any(indicator in cell_lower for indicator in ['name', 'email', 'phone', 'contact']):
                            # This might be a label cell, check adjacent cells for values
                            if cell_num + 1 < len(row.cells):
                                next_cell = row.cells[cell_num + 1].text.strip()
                                if next_cell:
                                    structured_data[cell_text] = next_cell
                                    print(f"DEBUG: Found structured data - {cell_text}: {next_cell}")
                
                # Add row as a single line if it has multiple meaningful cells
                if len(row_text) > 1:
                    combined_row = " | ".join(row_text)
                    text += combined_row + "\n"
                    print(f"DEBUG: Combined table row: '{combined_row}'")
                elif len(row_text) == 1:
                    text += row_text[0] + "\n"
            
            # Add structured data to main text with proper labels
            for label, value in structured_data.items():
                formatted_entry = f"{label}: {value}"
                text += formatted_entry + "\n"
                print(f"DEBUG: Added structured entry: '{formatted_entry}'")
            
            # Add the whole table text as well for fallback parsing
            if table_text.strip():
                text += "\n" + table_text.strip() + "\n"
        
        print(f"DEBUG: Total DOCX text extracted: {len(text)} characters")
        
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    
    return text

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read()
    except Exception as e:
        print(f"Error reading TXT: {e}")
    
    return text

def extract_email(text):
    """Extract email addresses from text, including hyperlinked emails"""
    print(f"DEBUG: Raw text length: {len(text)}")
    print(f"DEBUG: First 200 chars: {repr(text[:200])}")  # Show actual characters including special ones
    
    # Show all lines that contain common email indicators
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if any(indicator in line.lower() for indicator in ['@', 'mail', 'skype']):
            print(f"DEBUG: Line {i+1}: {repr(line)}")
    
    # Enhanced patterns to handle hyperlinked emails and various formats
    hyperlink_patterns = [
        # HTML-style links: <a href="mailto:email@domain.com">email@domain.com</a>
        r'<a\s+[^>]*href\s*=\s*["\']?mailto:([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})["\']?[^>]*>.*?</a>',
        # Markdown-style links: [email@domain.com](mailto:email@domain.com)
        r'\[([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\]\s*\(\s*mailto:[^)]+\)',
        # Simple mailto links: mailto:email@domain.com
        r'mailto\s*:\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        # Underlined or styled emails (often appear as _email@domain.com_ or similar)
        r'[_*]([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})[_*]',
        # Emails wrapped in angle brackets: <email@domain.com>
        r'<([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})>',
        # Emails with various separators/decorators
        r'[\[\(\{]([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})[\]\)\}]',
    ]
    
    # First, try to extract emails from hyperlink patterns
    for pattern in hyperlink_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
        if matches:
            print(f"DEBUG: Hyperlink pattern '{pattern}' found: {matches}")
            for match in matches:
                email_clean = match.strip().lower()
                if len(email_clean) > 5 and '.' in email_clean.split('@')[1]:
                    print(f"DEBUG: Returning hyperlinked email: {email_clean}")
                    return email_clean
    
    # Most aggressive email search - find ANY email pattern (with multiple attempts)
    # Try multiple regex patterns with different levels of strictness
    email_patterns = [
        r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',  # Standard pattern
        r'[a-zA-Z0-9._%+-]+\s*@\s*[a-zA-Z0-9.-]+\s*\.\s*[a-zA-Z]{2,}',  # With optional whitespace
        r'[a-zA-Z0-9._%-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,4}',  # Slightly different character set
    ]
    
    for pattern in email_patterns:
        all_emails = re.findall(pattern, text, re.IGNORECASE)
        print(f"DEBUG: Pattern '{pattern}' found emails: {all_emails}")
        
        if all_emails:
            # Clean and validate each email found
            for email in all_emails:
                # Remove any whitespace that might be in the email
                email_clean = re.sub(r'\s+', '', email).strip().lower()
                
                # Validate the email has proper structure
                if ('@' in email_clean and '.' in email_clean and 
                    len(email_clean) > 5 and 
                    len(email_clean.split('@')) == 2 and
                    '.' in email_clean.split('@')[1]):
                    print(f"DEBUG: Returning first valid email: {email_clean}")
                    return email_clean
    
    # If no emails found with simple pattern, try more specific searches
    # First, clean the text to handle encoded hyperlinks and special characters
    cleaned_text = text
    
    # Handle URL-encoded characters that might appear in hyperlinks
    import urllib.parse
    try:
        # Try to decode URL-encoded text (e.g., %40 for @, %2E for .)
        cleaned_text = urllib.parse.unquote(text)
        print(f"DEBUG: URL-decoded text found different content: {cleaned_text != text}")
    except:
        pass
    
    # Handle HTML entities
    import html
    try:
        # Decode HTML entities (e.g., &amp; for &, &#64; for @)
        html_decoded = html.unescape(cleaned_text)
        if html_decoded != cleaned_text:
            cleaned_text = html_decoded
            print(f"DEBUG: HTML-decoded text found different content")
    except:
        pass
    
    # Look for emails again in the cleaned text
    if cleaned_text != text:
        cleaned_emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', cleaned_text, re.IGNORECASE)
        print(f"DEBUG: Emails found in cleaned text: {cleaned_emails}")
        if cleaned_emails:
            for email in cleaned_emails:
                email_clean = email.strip().lower()
                if len(email_clean) > 5 and '.' in email_clean.split('@')[1]:
                    print(f"DEBUG: Returning email from cleaned text: {email_clean}")
                    return email_clean
    
    # Look for table-formatted contact information (common in structured resumes)
    table_patterns = [
        r'\|\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\s*\|',  # Email between table separators
        r'email\s*[\|:]\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',  # Email after "Email |" or "Email:"
        r'contact\s*[\|:]\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',  # Email after "Contact |"
        r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\s*\|',  # Email before table separator
    ]
    
    for pattern in table_patterns:
        matches = re.findall(pattern, cleaned_text, re.IGNORECASE)
        print(f"DEBUG: Table pattern '{pattern}' found: {matches}")
        if matches:
            email = matches[0].strip().lower()
            print(f"DEBUG: Returning table-formatted email: {email}")
            return email
    # Look for emails after specific labels (including enhanced hyperlink patterns)
    label_patterns = [
        # Standard label patterns
        r'mailto\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        r'mail\s*to\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        r'mailid\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        r'skypeid\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        r'skype\s*id\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        r'email\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        # Enhanced patterns for emails that might be linked or formatted
        r'e[\-\s]*mail\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        r'email\s*address\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        r'contact\s*email\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
    ]
    
    for pattern in label_patterns:
        matches = re.findall(pattern, cleaned_text, re.IGNORECASE)
        print(f"DEBUG: Pattern '{pattern}' found: {matches}")
        if matches:
            email = matches[0].strip().lower()
            print(f"DEBUG: Returning labeled email: {email}")
            return email
    
    # Manual line-by-line search for debugging (use both original and cleaned text)
    print("DEBUG: Manual line search...")
    
    # First try with cleaned text
    cleaned_lines = cleaned_text.split('\n')
    for i, line in enumerate(cleaned_lines):
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Check for email patterns in this line
        email_in_line = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', line_clean, re.IGNORECASE)
        if email_in_line:
            email = email_in_line.group(1).strip().lower()
            print(f"DEBUG: Found email in cleaned line {i+1}: '{email}' from line: '{line_clean}'")
            return email
    
    # If still not found, try with original text
    for i, line in enumerate(lines):
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Check for email patterns in this line
        email_in_line = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', line_clean, re.IGNORECASE)
        if email_in_line:
            email = email_in_line.group(1).strip().lower()
            print(f"DEBUG: Found email in original line {i+1}: '{email}' from line: '{line_clean}'")
            return email
    
    print("DEBUG: No email found anywhere")
    return "Not found"

def extract_phone(text):
    """Extract phone numbers from text"""
    phone_patterns = [
        r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b',
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
        r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',
        r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # International format
        r'Phone\s*:?\s*([+]?[\d\s\-\(\)\.]{10,})',  # After "Phone:" label
        r'Tel\s*:?\s*([+]?[\d\s\-\(\)\.]{10,})',    # After "Tel:" label
        r'Mobile\s*:?\s*([+]?[\d\s\-\(\)\.]{10,})',  # After "Mobile:" label
        r'Cell\s*:?\s*([+]?[\d\s\-\(\)\.]{10,})',   # After "Cell:" label
    ]
    
    # Clean text for better matching
    text_cleaned = text.replace('\n', ' ').replace('\r', ' ')
    
    for pattern in phone_patterns:
        phones = re.findall(pattern, text_cleaned, re.IGNORECASE)
        if phones:
            phone = phones[0]
            if isinstance(phone, tuple):
                # Format as (XXX) XXX-XXXX
                return f"({phone[0]}) {phone[1]}-{phone[2]}"
            else:
                # Clean and format the phone number
                cleaned = re.sub(r'[^\d+]', '', str(phone))
                if len(cleaned) >= 10:
                    if cleaned.startswith('+1'):
                        cleaned = cleaned[2:]
                    elif cleaned.startswith('1') and len(cleaned) == 11:
                        cleaned = cleaned[1:]
                    
                    if len(cleaned) == 10:
                        return f"({cleaned[:3]}) {cleaned[3:6]}-{cleaned[6:]}"
                    else:
                        return phone  # Return as-is if not standard format
    
    # Look for phone-like patterns line by line
    lines = text.split('\n')
    for line in lines:
        # Look for sequences of digits that could be phone numbers
        digits_only = re.sub(r'[^\d]', '', line)
        if len(digits_only) == 10:
            return f"({digits_only[:3]}) {digits_only[3:6]}-{digits_only[6:]}"
        elif len(digits_only) == 11 and digits_only.startswith('1'):
            cleaned = digits_only[1:]
            return f"({cleaned[:3]}) {cleaned[3:6]}-{cleaned[6:]}"
    
    return "Not found"

def calculate_name_score(name, line_number):
    """Calculate a score for how likely this is to be a real name"""
    score = 100  # Start with base score
    
    name_lower = name.lower()
    
    # HUGE bonus for line 1 - names are almost always on the first line
    if line_number == 0:  # line_number is 0-indexed
        score += 100
        print(f"DEBUG: Applied line 1 bonus (+100) to '{name}', new score: {score}")
    elif line_number == 1:  # Second line
        score += 50
        print(f"DEBUG: Applied line 2 bonus (+50) to '{name}', new score: {score}")
    elif line_number == 2:  # Third line
        score += 25
        print(f"DEBUG: Applied line 3 bonus (+25) to '{name}', new score: {score}")
    
    # Heavy penalties for section headers and common resume words
    section_words = [
        'summary', 'experience', 'skills', 'education', 'objective', 'profile',
        'background', 'history', 'qualifications', 'competencies', 'achievements',
        'professional', 'career', 'technical', 'personal', 'contact', 'information'
    ]
    
    # Additional penalties for location/address indicators
    location_words = [
        'crossing', 'republik', 'republic', 'township', 'complex', 'apartment', 'building',
        'block', 'sector', 'phase', 'extension', 'colony', 'society', 'enclave', 'park',
        'gardens', 'heights', 'plaza', 'mall', 'center', 'centre', 'residency', 'villa',
        'towers', 'manor', 'estate', 'homes', 'city', 'town', 'village', 'district',
        'nagar', 'vihar', 'puram', 'gram', 'pur', 'bad', 'garh', 'ganj', 'chowk'
    ]
    
    # Heavy penalties for job title indicators
    job_title_words = [
        'developer', 'engineer', 'manager', 'analyst', 'architect', 'consultant',
        'specialist', 'administrator', 'designer', 'tester', 'coordinator',
        'executive', 'associate', 'director', 'president', 'officer', 'lead',
        'senior', 'junior', 'principal', 'chief', 'head', 'supervisor',
        'technician', 'programmer', 'coder', 'operator', 'software', 'system',
        'business', 'technical', 'project', 'quality', 'database', 'network'
    ]
    
    penalty_applied = False
    for word in section_words:
        if word in name_lower:
            # Reduce penalty if it's line 1, as sometimes titles might appear there
            penalty = 30 if line_number == 0 else 50
            score -= penalty
            penalty_applied = True
            
    # Heavy penalty for location indicators
    for word in location_words:
        if word in name_lower:
            # Even location words get reduced penalty on line 1
            penalty = 40 if line_number == 0 else 70
            score -= penalty
            penalty_applied = True
            print(f"DEBUG: Applied location word penalty for '{word}' to '{name}', new score: {score}")
    
    # Extremely heavy penalty for job title indicators
    for word in job_title_words:
        if word in name_lower:
            # Reduced penalty for job title words on line 1 (sometimes names might contain these)
            penalty = 50 if line_number == 0 else 80
            score -= penalty
            penalty_applied = True
            print(f"DEBUG: Applied job title penalty for '{word}' to '{name}', new score: {score}")
            
    if penalty_applied:
        print(f"DEBUG: Applied penalties to '{name}', new score: {score}")
    
    # Penalize names with common section patterns
    if any(pattern in name_lower for pattern in ['summary', 'experience', 'skills']):
        penalty = 15 if line_number == 0 else 30
        score -= penalty
        print(f"DEBUG: Applied section pattern penalty to '{name}', new score: {score}")
    
    # Heavy penalty for names that sound like locations or addresses
    if any(location_term in name_lower for location_term in ['crossing republik', 'ajnara crossing', 'greater noida']):
        penalty = 60 if line_number == 0 else 100
        score -= penalty
        print(f"DEBUG: Applied obvious location penalty to '{name}', new score: {score}")
    
    # Extremely heavy penalty for obvious job titles
    job_title_patterns = [
        'senior software developer', 'software developer', 'senior developer',
        'software engineer', 'senior engineer', 'project manager', 'team leader',
        'business analyst', 'system analyst', 'technical architect'
    ]
    
    for pattern in job_title_patterns:
        if pattern in name_lower:
            penalty = 100 if line_number == 0 else 150
            score -= penalty
            print(f"DEBUG: Applied obvious job title penalty for '{pattern}' to '{name}', new score: {score}")
    
    # Earlier lines are more likely to contain names (but line 1 already got huge bonus)
    if line_number > 2:
        score -= (line_number - 2) * 5
    
    # Prefer 2-3 words over 4+ words
    word_count = len(name.split())
    if word_count == 2 or word_count == 3:
        score += 20
    elif word_count >= 4:
        score -= 10
    
    # Heavy penalty for very short or very long names
    if len(name) < 5:
        score -= 30
    elif len(name) > 30:
        score -= 20
    
    # Bonus for common name patterns
    words = name.split()
    if len(words) >= 2:
        # Check if first word looks like a first name (shorter, common patterns)
        if 3 <= len(words[0]) <= 10:
            score += 10
        # Check if last word looks like a last name
        if 3 <= len(words[-1]) <= 15:
            score += 10
    
    # Bonus for proper capitalization
    if all(word[0].isupper() for word in words):
        score += 15
    
    # Bonus for typical name characteristics
    if word_count == 2:  # First Last name pattern
        score += 15
    
    # Check if it looks like a person's name vs a title/section
    if all(len(word) >= 2 and word.isalpha() for word in words):
        score += 10
    
    # Additional validation: Check if words sound like real names vs locations/job titles
    # Penalty for words that are clearly not human names
    for word in words:
        word_lower = word.lower()
        if word_lower in ['ajnara', 'crossing', 'republik', 'republic', 'township', 'complex',
                         'developer', 'engineer', 'manager', 'senior', 'software', 'system']:
            penalty = 20 if line_number == 0 else 40
            score -= penalty
            print(f"DEBUG: Applied non-human-name penalty for word '{word}' to '{name}', new score: {score}")
    
    return score

def is_likely_person_name(name_candidate):
    """Check if a string is likely to be a person's name vs location/company/other"""
    name_lower = name_candidate.lower().strip()
    words = name_candidate.split()
    
    # Very obvious location/address patterns
    location_patterns = [
        'crossing republik', 'ajnara crossing', 'greater noida', 'sector', 'phase',
        'block', 'plot', 'flat', 'apartment', 'building', 'complex', 'tower',
        'mall', 'plaza', 'center', 'centre', 'city', 'town', 'village'
    ]
    
    for pattern in location_patterns:
        if pattern in name_lower:
            print(f"DEBUG: Rejected '{name_candidate}' - matches location pattern '{pattern}'")
            return False
    
    # Check for technology/technical terms
    technology_patterns = [
        'progress rdbms', 'progress database', 'progress openedge', 'openedge abl',
        'mysql database', 'oracle database', 'sql server', 'mongodb atlas',
        'microsoft office', 'visual studio', 'android studio', 'web development',
        'software development', 'mobile development', 'application development',
        'database management', 'project management', 'business intelligence',
        'data analysis', 'machine learning', 'artificial intelligence',
        'cloud computing', 'web services', 'api development', 'frontend development',
        'backend development', 'full stack', 'devops engineer', 'quality assurance',
        'technical skills', 'programming languages', 'software tools',
        'framework libraries', 'development tools', 'testing tools'
    ]
    
    for pattern in technology_patterns:
        if pattern in name_lower:
            print(f"DEBUG: Rejected '{name_candidate}' - matches technology pattern '{pattern}'")
            return False
    
    # Check for technical keywords
    technical_keywords = [
        'progress', 'rdbms', 'database', 'mysql', 'oracle', 'mongodb', 'redis',
        'javascript', 'typescript', 'python', 'java', 'angular', 'react', 'vue',
        'nodejs', 'express', 'django', 'flask', 'spring', 'laravel', 'symfony',
        'docker', 'kubernetes', 'jenkins', 'gitlab', 'github', 'azure', 'aws',
        'html', 'css', 'bootstrap', 'tailwind', 'sass', 'less', 'webpack',
        'framework', 'library', 'api', 'restful', 'soap', 'json', 'xml',
        'git', 'svn', 'agile', 'scrum', 'kanban', 'jira', 'confluence',
        'testing', 'selenium', 'cypress', 'jest', 'junit', 'postman',
        'linux', 'windows', 'ubuntu', 'centos', 'debian', 'fedora',
        'apache', 'nginx', 'tomcat', 'iis', 'server', 'hosting',
        'programming', 'coding', 'development', 'software', 'application',
        'web', 'mobile', 'desktop', 'frontend', 'backend', 'fullstack',
        'devops', 'cicd', 'deployment', 'automation', 'scripting'
    ]
    
    tech_keyword_count = 0
    for word in words:
        if word.lower() in technical_keywords:
            tech_keyword_count += 1
    
    # If any technical keywords are found, reject
    if tech_keyword_count > 0:
        print(f"DEBUG: Rejected '{name_candidate}' - contains technical keywords ({tech_keyword_count}/{len(words)})")
        return False
    
    # Check for job titles and professional designations
    job_title_patterns = [
        'senior software developer', 'software developer', 'software engineer',
        'senior developer', 'junior developer', 'lead developer', 'principal developer',
        'senior engineer', 'junior engineer', 'lead engineer', 'principal engineer',
        'project manager', 'senior manager', 'team leader', 'team lead',
        'business analyst', 'system analyst', 'data analyst', 'senior analyst',
        'architect', 'solution architect', 'technical architect', 'software architect',
        'consultant', 'senior consultant', 'technical consultant',
        'specialist', 'technical specialist', 'senior specialist',
        'administrator', 'system administrator', 'database administrator',
        'designer', 'ui designer', 'ux designer', 'graphic designer',
        'tester', 'qa tester', 'test engineer', 'quality analyst',
        'coordinator', 'project coordinator', 'technical coordinator',
        'executive', 'senior executive', 'business executive',
        'associate', 'senior associate', 'junior associate',
        'director', 'senior director', 'technical director',
        'vice president', 'assistant manager', 'deputy manager'
    ]
    
    for pattern in job_title_patterns:
        if pattern in name_lower:
            print(f"DEBUG: Rejected '{name_candidate}' - matches job title pattern '{pattern}'")
            return False
    
    # Check for job title keywords
    job_keywords = [
        'developer', 'engineer', 'manager', 'analyst', 'architect', 'consultant',
        'specialist', 'administrator', 'designer', 'tester', 'coordinator',
        'executive', 'associate', 'director', 'president', 'officer', 'lead',
        'senior', 'junior', 'principal', 'chief', 'head', 'supervisor',
        'technician', 'programmer', 'coder', 'administrator', 'operator'
    ]
    
    job_keyword_count = 0
    for word in words:
        if word.lower() in job_keywords:
            job_keyword_count += 1
    
    # If more than half the words are job title keywords, reject
    if len(words) > 0 and job_keyword_count >= len(words) / 2:
        print(f"DEBUG: Rejected '{name_candidate}' - too many job title words ({job_keyword_count}/{len(words)})")
        return False
    
    # Check for email-related patterns
    email_patterns = [
        'e-mail', 'email', 'mail', 'gmail', 'outlook', 'yahoo', 'hotmail',
        'contact', 'phone', 'mobile', 'cell', 'tel', 'fax', 'address',
        'linkedin', 'skype', 'website', 'blog', 'portfolio'
    ]
    
    for pattern in email_patterns:
        if pattern in name_lower:
            print(f"DEBUG: Rejected '{name_candidate}' - contains email/contact pattern '{pattern}'")
            return False
    
    # Check for patterns with colons (common in contact info like "E-Mail:", "Phone:")
    if ':' in name_candidate:
        print(f"DEBUG: Rejected '{name_candidate}' - contains colon (likely a label)")
        return False
    
    # Check for @ symbol or email-like patterns
    if '@' in name_candidate or any(char.isdigit() for char in name_candidate):
        print(f"DEBUG: Rejected '{name_candidate}' - contains @ symbol or digits")
        return False
    
    # Check for location-specific words
    location_words = [
        'crossing', 'republik', 'republic', 'township', 'residency', 'gardens',
        'heights', 'manor', 'estate', 'homes', 'enclave', 'colony', 'society',
        'nagar', 'vihar', 'puram', 'gram', 'garh', 'ganj', 'chowk', 'marg'
    ]
    
    location_word_count = 0
    for word in words:
        if word.lower() in location_words:
            location_word_count += 1
    
    # If more than half the words are location indicators, reject
    if len(words) > 0 and location_word_count >= len(words) / 2:
        print(f"DEBUG: Rejected '{name_candidate}' - too many location words ({location_word_count}/{len(words)})")
        return False
    
    # Check if it has characteristics of a human name
    if len(words) == 2 or len(words) == 3:
        # Check if words look like typical names (not too long, alphabetic)
        for word in words:
            # Names typically don't have very long words
            if len(word) > 15:
                print(f"DEBUG: Rejected '{name_candidate}' - word '{word}' too long for a name")
                return False
            
            # Check for patterns that suggest it's not a name
            word_lower = word.lower()
            if word_lower in ['crossing', 'republik', 'republic', 'ajnara', 'township', 'complex',
                             'e-mail', 'email', 'mail', 'phone', 'mobile', 'contact', 'address',
                             'developer', 'engineer', 'manager', 'analyst', 'architect', 'consultant',
                             'senior', 'junior', 'lead', 'principal', 'software', 'system']:
                print(f"DEBUG: Rejected '{name_candidate}' - contains obvious non-name word '{word}'")
                return False
    
    # Additional checks for Indian context
    # Some patterns that are definitely not names
    indian_location_patterns = [
        'crossing republik', 'crossing republic', 'greater noida', 'new delhi',
        'bangalore', 'mumbai', 'chennai', 'hyderabad', 'pune', 'kolkata'
    ]
    
    for pattern in indian_location_patterns:
        if pattern in name_lower:
            print(f"DEBUG: Rejected '{name_candidate}' - matches Indian location pattern '{pattern}'")
            return False
    
    print(f"DEBUG: '{name_candidate}' passed person name validation")
    return True

def clean_name_candidate(name_candidate):
    """Clean a name candidate by removing email/contact labels and extra text"""
    # Remove common patterns that get attached to names
    patterns_to_remove = [
        r'\s*e-?mail\s*:?\s*$',  # Remove "e-mail:" or "email:" at the end
        r'\s*phone\s*:?\s*$',    # Remove "phone:" at the end
        r'\s*mobile\s*:?\s*$',   # Remove "mobile:" at the end
        r'\s*contact\s*:?\s*$',  # Remove "contact:" at the end
        r'\s*tel\s*:?\s*$',      # Remove "tel:" at the end
        r'\s*cell\s*:?\s*$',     # Remove "cell:" at the end
        r'\s*address\s*:?\s*$',  # Remove "address:" at the end
    ]
    
    cleaned_name = name_candidate.strip()
    
    for pattern in patterns_to_remove:
        cleaned_name = re.sub(pattern, '', cleaned_name, flags=re.IGNORECASE).strip()
    
    # Remove trailing punctuation and whitespace
    cleaned_name = re.sub(r'[:\-,\s]+$', '', cleaned_name).strip()
    
    # If the cleaning removed too much, return the original
    if len(cleaned_name) < 3:
        return name_candidate
    
    # Check if the cleaned name still has valid words
    words = cleaned_name.split()
    if len(words) >= 2 and all(len(word) >= 2 and word.replace('.', '').isalpha() for word in words):
        print(f"DEBUG: Cleaned name from '{name_candidate}' to '{cleaned_name}'")
        return cleaned_name
    else:
        print(f"DEBUG: Cleaning would make name invalid, keeping original: '{name_candidate}'")
        return name_candidate

def extract_name(text):
    """Extract name from text with enhanced header detection"""
    print(f"DEBUG: Name extraction from text length: {len(text)}")
    
    lines = text.split('\n')
    print(f"DEBUG: First 15 lines for name extraction:")
    for i, line in enumerate(lines[:15]):
        if line.strip():
            print(f"  Line {i+1}: '{line.strip()}'")
    
    # Common words that indicate it's NOT a name
    non_name_indicators = [
        'resume', 'cv', 'curriculum', 'vitae', 'curriculam', 'profile', 'summary', 'objective',
        'email', 'phone', 'address', 'contact', 'linkedin', 'github', 'www', 'http',
        'experience', 'education', 'skills', 'projects', 'work', 'employment',
        'mailto', 'skype', 'gmail', 'outlook', 'yahoo', 'hotmail', '.com', '.org',
        'mobile', 'cell', 'tel', 'fax', 'website', 'blog', 'portfolio',
        # Document type headers (common resume/CV headers)
        'curriculum vitae', 'curriculam vitae', 'resume', 'biodata', 'bio data',
        'personal details', 'personal information', 'contact details',
        # Section headers that are definitely not names
        'professional summary', 'career summary', 'executive summary',
        'work experience', 'professional experience', 'employment history',
        'technical skills', 'core competencies', 'key qualifications',
        'educational background', 'academic background', 'certifications',
        'achievements', 'accomplishments', 'awards', 'honors',
        'personal information', 'contact information', 'references',
        'career objective', 'professional objective', 'personal profile', 'technical expertise',
        'experience summary',
        # Job titles and professional designations
        'software developer', 'senior software developer', 'junior software developer',
        'software engineer', 'senior software engineer', 'junior software engineer',
        'web developer', 'full stack developer', 'frontend developer', 'backend developer',
        'mobile developer', 'application developer', 'systems developer',
        'project manager', 'senior project manager', 'assistant project manager',
        'program manager', 'product manager', 'business analyst', 'system analyst',
        'data analyst', 'senior analyst', 'junior analyst', 'technical analyst',
        'solution architect', 'software architect', 'system architect', 'enterprise architect',
        'technical architect', 'lead architect', 'principal architect',
        'team leader', 'team lead', 'tech lead', 'technical lead', 'lead developer',
        'senior developer', 'junior developer', 'principal developer',
        'consultant', 'senior consultant', 'technical consultant', 'it consultant',
        'specialist', 'technical specialist', 'senior specialist', 'subject matter expert',
        'administrator', 'system administrator', 'database administrator', 'network administrator',
        'designer', 'ui designer', 'ux designer', 'graphic designer', 'web designer',
        'quality assurance', 'qa engineer', 'test engineer', 'qa analyst', 'tester',
        'coordinator', 'project coordinator', 'technical coordinator', 'program coordinator',
        'executive', 'senior executive', 'business executive', 'account executive',
        'associate', 'senior associate', 'junior associate', 'business associate',
        'director', 'senior director', 'technical director', 'managing director',
        'vice president', 'assistant manager', 'deputy manager', 'general manager',
        # Common job level indicators
        'senior', 'junior', 'lead', 'principal', 'chief', 'head', 'assistant', 'deputy',
        # Location and address indicators
        'crossing', 'republik', 'republic', 'township', 'complex', 'apartment', 'building',
        'block', 'sector', 'phase', 'extension', 'colony', 'society', 'enclave', 'park',
        'gardens', 'heights', 'plaza', 'mall', 'center', 'centre', 'residency', 'villa',
        'towers', 'manor', 'estate', 'homes', 'city', 'town', 'village', 'district',
        'state', 'country', 'street', 'road', 'avenue', 'lane', 'nagar', 'vihar',
        'puram', 'gram', 'pur', 'bad', 'garh', 'ganj', 'chowk', 'marg', 'path',
        # Indian location specific terms
        'delhi', 'mumbai', 'bangalore', 'chennai', 'kolkata', 'hyderabad', 'pune',
        'ahmedabad', 'gurgaon', 'noida', 'faridabad', 'ghaziabad', 'greater noida',
        # Company/Organization indicators
        'ltd', 'limited', 'pvt', 'private', 'company', 'corporation', 'corp', 'inc',
        'incorporated', 'llc', 'enterprises', 'solutions', 'services', 'technologies',
        'systems', 'consultancy', 'consulting', 'group', 'organization', 'institution'
    ]
    
    # Section header patterns (common resume section names)
    section_patterns = [
        r'^(professional|career|executive|personal)\s+(summary|profile|objective)',
        r'^(work|professional|employment)\s+experience',
        r'^(technical|core|key)\s+(skills|competencies|qualifications)',
        r'^(educational|academic)\s+background',
        r'^(contact|personal)\s+information',
        r'^experience\s+(summary|overview)',
        r'^summary\s+of\s+(qualifications|experience)',
        r'^\w+\s+(summary|experience|skills|education|background|information)$'
    ]
    
    # Strategy 1: Look for the most likely name in first few lines
    candidate_names = []
    
    for i, line in enumerate(lines[:10]):  # Check first 10 lines
        line = line.strip()
        #if not line or len(line) < 3:
        #  continue
        print(f"DEBUG: line {i+1} : '{line}' len(line): '{len(line)}'")

        # Skip lines with obvious non-name content
        line_lower = line.lower()
        
        # Specific check for document type headers (more strict)
        document_headers = [
            'curriculum vitae', 'curriculam vitae', 'resume', 'cv', 'biodata', 'bio data'
        ]
        
        # Check if the entire line (or most of it) is a document header
        line_words = line_lower.split()
        is_document_header = False
        for header in document_headers:
            header_words = header.split()
            if len(line_words) <= 3 and all(word in line_lower for word in header_words):
                print(f"DEBUG: Skipping line {i+1} (document header): '{line}'")
                is_document_header = True
                break
        
        if is_document_header:
            continue
        
        if any(indicator in line_lower for indicator in non_name_indicators):
            print(f"DEBUG: Skipping line {i+1} (contains non-name indicator): '{line}'")
            continue
        
        # Check if line matches section header patterns
        is_section_header = False
        for pattern in section_patterns:
            if re.match(pattern, line_lower):
                print(f"DEBUG: Skipping line {i+1} (matches section pattern '{pattern}'): '{line}'")
                is_section_header = True
                break
        
        if is_section_header:
            continue
            
        # Skip lines with numbers, @ symbols, or too much punctuation
        if '@' in line or any(char.isdigit() for char in line):
            print(f"DEBUG: Skipping line {i+1} (contains @ or digits): '{line}'")
            continue
        
        # Skip lines that are too long to be names (likely descriptions)
        if len(line) > 40:
            print(f"DEBUG: Skipping line {i+1} (too long for a name): '{line}'")
            continue
        
        # Special handling for table-formatted names (common pattern: Name | Email | Phone)
        if '|' in line:
            table_parts = [part.strip() for part in line.split('|')]
            print(f"DEBUG: Found table-formatted line {i+1}: {table_parts}")
            
            # Look for name in first column (most common pattern)
            if len(table_parts) > 0:
                potential_name = table_parts[0].strip()
                
                # Clean the potential name first
                cleaned_name = clean_name_candidate(potential_name)
                
                # Validate this looks like a name
                name_words = cleaned_name.split()
                if (2 <= len(name_words) <= 4 and 
                    all(len(word) >= 2 and word.replace('.', '').replace("'", '').isalpha() 
                        for word in name_words) and
                    not any(indicator in cleaned_name.lower() for indicator in non_name_indicators) and
                    is_likely_person_name(cleaned_name)):
                    
                    score = calculate_name_score(cleaned_name, i)
                    
                    # Check for duplicates before adding
                    is_duplicate = False
                    for existing_name, _, _ in candidate_names:
                        if existing_name.lower() == cleaned_name.lower():
                            is_duplicate = True
                            print(f"DEBUG: Skipping duplicate table candidate '{cleaned_name}' from line {i+1}")
                            break
                    
                    if not is_duplicate:
                        candidate_names.append((cleaned_name, score + 25, i))  # Bonus for table format
                        print(f"DEBUG: Table name candidate from line {i+1}: '{cleaned_name}' (score: {score + 25})")
                else:
                    print(f"DEBUG: Rejected table candidate from line {i+1}: '{cleaned_name}' (failed validation)")
            
            # Also check other columns for names (less common but possible)
            for col_num, part in enumerate(table_parts[1:], 1):
                if (not '@' in part and not any(char.isdigit() for char in part) and
                    len(part.split()) >= 2 and len(part) < 30):
                    
                    # Clean the potential name first
                    cleaned_part = clean_name_candidate(part)
                    part_words = cleaned_part.split()
                    
                    if (all(len(word) >= 2 and word.replace('.', '').replace("'", '').isalpha() 
                           for word in part_words) and is_likely_person_name(cleaned_part)):
                        score = calculate_name_score(cleaned_part, i)
                        
                        # Check for duplicates before adding
                        is_duplicate = False
                        for existing_name, _, _ in candidate_names:
                            if existing_name.lower() == cleaned_part.lower():
                                is_duplicate = True
                                print(f"DEBUG: Skipping duplicate table candidate '{cleaned_part}' from column {col_num + 1}, line {i+1}")
                                break
                        
                        if not is_duplicate:
                            candidate_names.append((cleaned_part, score + 10, i))  # Smaller bonus for non-first column
                            print(f"DEBUG: Table name candidate from column {col_num + 1}, line {i+1}: '{cleaned_part}' (score: {score + 10})")
                    else:
                        print(f"DEBUG: Rejected table candidate from column {col_num + 1}, line {i+1}: '{cleaned_part}' (failed validation)")
            
            continue  # Skip normal processing for table lines
            
        # Clean the line of common formatting
        clean_line = line.replace('|', ' ').replace('•', ' ').replace('_', ' ')
        clean_line = re.sub(r'[^\w\s\'-.]', ' ', clean_line)  # Keep only letters, spaces, apostrophes, hyphens, dots
        clean_line = re.sub(r'\s+', ' ', clean_line).strip()
        
        if not clean_line:
            continue
            
        words = clean_line.split()
        
        # Name validation: 2-4 words, each word 2+ characters, mostly alphabetic
        if 2 <= len(words) <= 4:
            valid_words = []
            for word in words:
                # Clean word (remove dots, check if it's alphabetic)
                clean_word = word.replace('.', '').replace("'", '')
                if len(clean_word) >= 2 and clean_word.isalpha():
                    # Check if word looks like a name part (capitalize first letter)
                    if clean_word[0].isupper() or word.istitle():
                        valid_words.append(word.title())  # Ensure proper capitalization
                    
            if len(valid_words) >= 2:  # At least 2 valid name parts
                potential_name = ' '.join(valid_words)
                
                # Clean the potential name first
                cleaned_name = clean_name_candidate(potential_name)
                
                # Add validation check for person name
                if is_likely_person_name(cleaned_name):
                    score = calculate_name_score(cleaned_name, i)
                    
                    # Check for duplicates before adding
                    is_duplicate = False
                    for existing_name, _, _ in candidate_names:
                        if existing_name.lower() == cleaned_name.lower():
                            is_duplicate = True
                            print(f"DEBUG: Skipping duplicate candidate '{cleaned_name}' from line {i+1}")
                            break
                    
                    if not is_duplicate:
                        candidate_names.append((cleaned_name, score, i))
                        print(f"DEBUG: Candidate name from line {i+1}: '{cleaned_name}' (score: {score})")
                else:
                    print(f"DEBUG: Rejected candidate from line {i+1}: '{cleaned_name}' (failed person name validation)")
    
    # Strategy 2: If no good candidates, look for any capitalized sequences
    if not candidate_names:
        print("DEBUG: No candidates found, trying broader search...")
        for i, line in enumerate(lines[:15]):
            line = line.strip()
            if not line:
                continue
                
            # Find sequences of capitalized words
            words = line.split()
            cap_sequence = []
            
            for word in words:
                clean_word = re.sub(r'[^\w]', '', word)  # Remove all punctuation
                if (len(clean_word) >= 2 and clean_word[0].isupper() and 
                    clean_word.isalpha() and clean_word not in non_name_indicators):
                    cap_sequence.append(word.title())
                else:
                    if len(cap_sequence) >= 2:  # End of sequence, check if it's a good name
                        potential_name = ' '.join(cap_sequence)
                        cleaned_name = clean_name_candidate(potential_name)
                        if is_likely_person_name(cleaned_name):
                            score = calculate_name_score(cleaned_name, i)
                            
                            # Check for duplicates before adding
                            is_duplicate = False
                            for existing_name, _, _ in candidate_names:
                                if existing_name.lower() == cleaned_name.lower():
                                    is_duplicate = True
                                    print(f"DEBUG: Skipping duplicate capitalized sequence '{cleaned_name}' from line {i+1}")
                                    break
                            
                            if not is_duplicate:
                                candidate_names.append((cleaned_name, score, i))
                                print(f"DEBUG: Capitalized sequence from line {i+1}: '{cleaned_name}' (score: {score})")
                        else:
                            print(f"DEBUG: Rejected capitalized sequence from line {i+1}: '{cleaned_name}' (failed person name validation)")
                    cap_sequence = []
            
            # Check final sequence
            if len(cap_sequence) >= 2:
                potential_name = ' '.join(cap_sequence)
                cleaned_name = clean_name_candidate(potential_name)
                if is_likely_person_name(cleaned_name):
                    score = calculate_name_score(cleaned_name, i)
                    
                    # Check for duplicates before adding
                    is_duplicate = False
                    for existing_name, _, _ in candidate_names:
                        if existing_name.lower() == cleaned_name.lower():
                            is_duplicate = True
                            print(f"DEBUG: Skipping duplicate final sequence '{cleaned_name}' from line {i+1}")
                            break
                    
                    if not is_duplicate:
                        candidate_names.append((cleaned_name, score, i))
                        print(f"DEBUG: Final capitalized sequence from line {i+1}: '{cleaned_name}' (score: {score})")
                else:
                    print(f"DEBUG: Rejected final capitalized sequence from line {i+1}: '{cleaned_name}' (failed person name validation)")
    
    # Select the best candidate name
    if candidate_names:
        # Sort by score (higher is better)
        candidate_names.sort(key=lambda x: x[1], reverse=True)
        best_name = candidate_names[0][0]
        print(f"DEBUG: Selected best name: '{best_name}' from {len(candidate_names)} candidates")
        return best_name
    
    print("DEBUG: No valid name found")
    return "Not found"

def load_education_keywords():
    """Load education keywords from external file"""
    keywords_file = os.path.join(os.path.dirname(__file__), 'keywords', 'education_keywords.txt')
    education_keywords = []
    
    try:
        with open(keywords_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                # Skip empty lines and comments
                if line and not line.startswith('#'):
                    education_keywords.append(line.lower())
                    
    except FileNotFoundError:
        print("WARNING: education_keywords.txt not found, using fallback keywords")
        # Fallback to basic keywords if file not found
        education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'diploma', 'degree',
            'b.tech', 'b.e', 'b.sc', 'm.sc', 'mba', 'bca', 'mca',
            '10th', '12th', 'sslc', 'hsc', 'graduation', 'education'
        ]
    
    return education_keywords

def extract_education(text):
    import re  # Make sure re is imported
    
    print("DEBUG: Starting education extraction...")
    
    # Load education keywords from external file
    education_keywords = load_education_keywords()
    
    print(f"DEBUG: Loaded {len(education_keywords)} education keywords")
    
    # Keywords that should NOT be considered educational (EXCLUDE these)
    exclude_keywords = [
        'experience', 'work', 'job', 'employment', 'company', 'project', 'skill',
        'software', 'programming', 'development', 'management',
        'responsibility', 'achievement', 'award', 'training', 'workshop',
        'seminar', 'conference', 'publication', 'research', 'objective',
        'summary', 'profile', 'career', 'professional', 'expertise', 'knowledge',
        'tools', 'framework', 'database', 'language', 'platform', 'environment', 
        'private', 'limited', 'technologies', 'solutions', 'services', 'systems'
    ]
    
    # University/institution keywords that should not be treated as degrees
    institution_keywords = [
        'university', 'college', 'institute', 'school', 'academy', 'campus',
        'board', 'cbse', 'icse', 'state board', 'central board'
    ]
    
    lines = text.split('\n')
    education_info = []
    
    print(f"DEBUG: Processing {len(lines)} lines for education...")
    
    for i, line in enumerate(lines):
        line_clean = line.strip()
        line_lower = line_clean.lower()
        
        # Skip empty lines or very short lines
        if len(line_clean) < 5:
            continue
            
        # FIRST: Check if line contains exclude keywords (reject immediately)
        exclude_found = [word for word in exclude_keywords if word in line_lower]
        if exclude_found:
            print(f"DEBUG: Line {i+1} EXCLUDED (contains {exclude_found}): '{line_clean[:50]}...'")
            continue
        
        # SECOND: Check if line contains education keywords (using word boundaries for better matching)
        matched_keywords = []
        for keyword in education_keywords:
            # Use word boundaries for short keywords to avoid substring matches
            if len(keyword) <= 4:
                pattern = r'\b' + re.escape(keyword) + r'\b'
                if re.search(pattern, line_lower):
                    matched_keywords.append(keyword)
            else:
                # For longer keywords, use simple substring matching
                if keyword in line_lower:
                    matched_keywords.append(keyword)
        
        if matched_keywords:
            print(f"DEBUG: Line {i+1} MATCHED keywords {matched_keywords}: '{line_clean}'")
            
            # Check if this is just an institution name (not a degree)
            is_just_institution = False
            if any(inst in line_lower for inst in institution_keywords):
                # If it only contains institution keywords and no degree keywords
                degree_keywords = ['bachelor', 'master', 'diploma', 'b.tech', 'b.e', 'b.sc', 'm.sc', 
                                 'mba', 'phd', 'bca', 'mca', '10th', '12th', 'sslc', 'hsc', 'degree']
                if not any(deg in line_lower for deg in degree_keywords):
                    # Additional check: if line is very short and mostly institution name
                    clean_line_words = line_clean.split()
                    if len(clean_line_words) <= 4 and any(inst in line_lower for inst in ['university', 'college', 'institute']):
                        is_just_institution = True
                        print(f"DEBUG: Line {i+1} REJECTED (just institution name): '{line_clean}'")
                        continue
            
            # Additional validation - check if this looks like a real education entry
            is_valid_education = False
            
            # Must contain at least one strong education indicator
            strong_indicators = [
                'degree', 'bachelor', 'master', 'diploma', 'certificate', 'graduation',
                'b.tech', 'b.e', 'b.sc', 'm.sc', 'mba', 'phd', 'bca', 'mca', 
                '10th', '12th', 'sslc', 'hsc', 'class 10', 'class 12'
            ]
            
            if any(indicator in line_lower for indicator in strong_indicators):
                is_valid_education = True
            
            # OR section headers with years
            elif any(edu_word in line_lower for edu_word in ['education', 'qualification', 'academic']):
                if re.search(r'\b(19|20)\d{2}\b', line_clean):
                    is_valid_education = True
            
            if not is_valid_education:
                print(f"DEBUG: Line {i+1} REJECTED (not strong enough education indicator): '{line_clean}'")
                continue
            
            # Clean up qualification name
            qualification = line_clean
            
            # Remove common prefixes/suffixes
            for prefix in ['•', '-', '*', ':', '▪', '○', '►', '→']:
                if qualification.startswith(prefix):
                    qualification = qualification[1:].strip()
            
            # Remove common section headers
            section_headers = ['education:', 'qualification:', 'academic:', 'degrees:', 'qualifications:']
            for header in section_headers:
                if qualification.lower().startswith(header):
                    qualification = qualification[len(header):].strip()
            
            # Skip if qualification is too generic or empty
            if (len(qualification) < 8 or 
                qualification.lower() in ['education', 'qualification', 'academic', 'qualifications']):
                print(f"DEBUG: Line {i+1} SKIPPED (too generic): '{qualification}'")
                continue
            
            # Check if the line already contains structured information (year and grade/percentage)
            # If it does, use it as-is without extracting separately
            contains_year = bool(re.search(r'\b(19|20)\d{2}\b', qualification))
            contains_percentage = bool(re.search(r'\d+(?:\.\d+)?\s*%|\d+(?:\.\d+)?\s*(?:gpa|cgpa)|first\s+class|second\s+class|distinction', qualification.lower()))
            
            if contains_year and contains_percentage:
                # Line already contains complete information
                education_entry = qualification
                print(f"DEBUG: Using complete line as-is: {education_entry}")
            else:
                # Extract missing information from nearby lines
                year = "Not specified"
                percentage = "Not specified"
                
                # Search only in current line first, then next 1 line to avoid cross-contamination
                search_text = qualification
                
                # First try to extract from current line only
                if not contains_year:
                    year_matches = re.findall(r'\b(19|20)(\d{2})\b', qualification)
                    if year_matches:
                        year = year_matches[0][0] + year_matches[0][1]
                        print(f"DEBUG: Found year in current line: {year}")
                
                if not contains_percentage:
                    percentage_patterns = [
                        r'(\d{1,3}(?:\.\d{1,2})?)\s*%',
                        r'(\d{1,3}(?:\.\d{1,2})?)\s*percent',
                        r'(\d\.\d+)\s*(?:gpa|cgpa)',
                        r'grade\s*[:\-]?\s*([a-f][\+\-]?)',
                        r'(first\s+class|second\s+class|third\s+class)',
                        r'(distinction|merit|pass)',
                        r'(\d{1,3})\s*marks?'
                    ]
                    
                    for pattern in percentage_patterns:
                        matches = re.findall(pattern, qualification, re.IGNORECASE)
                        if matches:
                            if isinstance(matches[0], tuple):
                                non_empty_parts = [str(part) for part in matches[0] if part]
                                if non_empty_parts:
                                    percentage = ' '.join(non_empty_parts).strip()
                            else:
                                percentage = str(matches[0]).strip()
                            print(f"DEBUG: Found percentage/grade in current line: {percentage}")
                            break
                
                # If still not found, try next line only (and only if it's related)
                if (year == "Not specified" or percentage == "Not specified") and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if (len(next_line) > 3 and 
                        not any(exclude in next_line.lower() for exclude in exclude_keywords) and
                        not any(keyword in next_line.lower() for keyword in education_keywords[:20])):  # Avoid picking up other education entries
                        
                        if year == "Not specified":
                            year_matches = re.findall(r'\b(19|20)(\d{2})\b', next_line)
                            if year_matches:
                                year = year_matches[0][0] + year_matches[0][1]
                                print(f"DEBUG: Found year in next line: {year}")
                        
                        if percentage == "Not specified":
                            for pattern in percentage_patterns:
                                matches = re.findall(pattern, next_line, re.IGNORECASE)
                                if matches:
                                    if isinstance(matches[0], tuple):
                                        non_empty_parts = [str(part) for part in matches[0] if part]
                                        if non_empty_parts:
                                            percentage = ' '.join(non_empty_parts).strip()
                                    else:
                                        percentage = str(matches[0]).strip()
                                    print(f"DEBUG: Found percentage/grade in next line: {percentage}")
                                    break
                
                # Format the education entry
                if year != "Not specified" or percentage != "Not specified":
                    education_entry = f"{qualification} - {year} - {percentage}"
                else:
                    education_entry = qualification
            
            print(f"DEBUG: Formatted entry: {education_entry}")
            
            # Avoid duplicates and overly long entries
            if education_entry not in education_info and len(qualification) < 200:
                education_info.append(education_entry)
                print(f"DEBUG: ADDED education entry: {education_entry}")
            else:
                print(f"DEBUG: DUPLICATE or TOO LONG, skipping: {education_entry}")
    
    print(f"DEBUG: Total education entries found: {len(education_info)}")
    
    # Remove duplicates while preserving order
    seen = set()
    unique_education = []
    for edu in education_info:
        if edu not in seen:
            seen.add(edu)
            unique_education.append(edu)
    
    print(f"DEBUG: Final education list: {unique_education}")
    return unique_education[:10] if unique_education else ["Not found"]

def load_skills_keywords():
    """Load skills keywords from external files"""
    keywords_dir = os.path.join(os.path.dirname(__file__), 'keywords')
    
    technical_skills = []
    functional_skills = []
    domain_skills = []
    
    # Load technical skills
    try:
        with open(os.path.join(keywords_dir, 'technical_skills.txt'), 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):
                    technical_skills.append(line.lower())
    except FileNotFoundError:
        print("WARNING: technical_skills.txt not found, using fallback")
        technical_skills = ['python', 'java', 'javascript', 'sql', 'html', 'css']
    
    # Load functional skills
    try:
        with open(os.path.join(keywords_dir, 'functional_skills.txt'), 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):
                    functional_skills.append(line.lower())
    except FileNotFoundError:
        print("WARNING: functional_skills.txt not found, using fallback")
        functional_skills = ['project management', 'team leadership', 'communication']
    
    # Load domain skills
    try:
        with open(os.path.join(keywords_dir, 'domain_skills.txt'), 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):
                    domain_skills.append(line.lower())
    except FileNotFoundError:
        print("WARNING: domain_skills.txt not found, using fallback")
        domain_skills = ['healthcare', 'finance', 'manufacturing', 'retail']
    
    return technical_skills, functional_skills, domain_skills

def extract_skills(text):
    """Extract skills from text and categorize into technical, functional, and domain skills"""
    
    # Load skills keywords from external files
    technical_skills, functional_skills, domain_skills = load_skills_keywords()
    
    print(f"DEBUG: Loaded {len(technical_skills)} technical, {len(functional_skills)} functional, and {len(domain_skills)} domain skills")
    
    # Keywords that should NOT be considered Skill (EXCLUDE these)
    exclude_keywords = [
        'private', 'limited', 'confidential', 'proprietary', 'company', 'organization', 'india',
        'corporation', 'incorporated', 'ltd', 'pvt', 'pvt ltd', 'private limited',
        'technologies', 'solutions', 'services', 'systems', 'consulting', 'consultancy', 'group', 'enterprise', 'group of companies',
        'technologies', 'technologies pvt ltd', 'technologies limited', 'technologies private limited', 'since', 'established', 'founded',
        'years of experience', 'experience in', 'expertise in', 'knowledge of', 'companies', 'duration', 'operating system'
    ]
    text_lower = text.lower()
    
    # Found skills categorized
    found_technical = []
    found_functional = []
    found_domain = []
    
    # Extract technical skills
    for skill in technical_skills:
        if skill.lower() in text_lower:
            found_technical.append(skill.title())
    
    # Extract functional skills
    for skill in functional_skills:
        if skill.lower() in text_lower:
            found_functional.append(skill.title())
    
    # Extract domain skills
    for skill in domain_skills:
        if skill.lower() in text_lower:
            found_domain.append(skill.title())
    
    # Look for skills section specifically for additional parsing
    lines = text.split('\n')
    
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in ['skill', 'technical', 'competenc', 'expert']):
            # Get next few lines after skills header
            for j in range(i + 1, min(i + 8, len(lines))):
                skills_line = lines[j].strip()

                # FIRST: Check if line contains exclude keywords (reject immediately)
                exclude_found = [word for word in exclude_keywords if word in skills_line.lower()]
                if exclude_found:
                    print(f"DEBUG: Line {i+1} EXCLUDED (contains {exclude_found}): '{skills_line[:50]}...'")
                    continue
                
                if skills_line and not any(section in skills_line.lower() for section in 
                                         ['experience', 'education', 'work', 'employment', 'project']):
                    
                    # Parse comma-separated or bullet-pointed skills
                    if ',' in skills_line:
                        line_skills = [s.strip() for s in skills_line.split(',') if len(s.strip()) > 2]
                        
                        # Categorize these additional skills
                        for skill in line_skills:
                            skill_lower = skill.lower()
                            
                            # Check if it's a technical skill
                            if any(tech in skill_lower for tech in ['programming', 'development', 'framework', 
                                                                    'database', 'software', 'tool', 'technology',
                                                                    'language', 'platform', 'system', 'api']):
                                if skill not in found_technical:
                                    found_technical.append(skill)
                            
                            # Check if it's a functional skill
                            elif any(func in skill_lower for func in ['management', 'analysis', 'planning',
                                                                      'leadership', 'communication', 'process',
                                                                      'methodology', 'testing', 'design']):
                                if skill not in found_functional:
                                    found_functional.append(skill)
                            
                            # Check if it's a domain skill
                            elif any(domain in skill_lower for domain in ['business', 'industry', 'domain',
                                                                          'sector', 'finance', 'healthcare',
                                                                          'retail', 'manufacturing']):
                                if skill not in found_domain:
                                    found_domain.append(skill)
                            
                            # If uncategorized but seems like a skill, add to technical by default
                            elif len(skill) > 2 and skill not in found_technical:
                                found_technical.append(skill)
                    
                    elif any(bullet in skills_line for bullet in ['•', '*', '-', '►']):
                        skill = skills_line
                        for bullet in ['•', '*', '-', '►']:
                            skill = skill.replace(bullet, '').strip()
                        
                        if skill and len(skill) > 2:
                            # Categorize bullet point skills using same logic
                            skill_lower = skill.lower()
                            if any(tech in skill_lower for tech in ['programming', 'development', 'framework']):
                                found_technical.append(skill)
                            elif any(func in skill_lower for func in ['management', 'analysis', 'leadership']):
                                found_functional.append(skill)
                            else:
                                found_technical.append(skill)  # Default to technical
    
    # Remove duplicates while preserving order
    found_technical = list(dict.fromkeys(found_technical))
    found_functional = list(dict.fromkeys(found_functional))
    found_domain = list(dict.fromkeys(found_domain))
    
    # Create structured skills table
    skills_table = {
        'Technical Skills': found_technical[:25] if found_technical else ["None identified"],
        'Functional Skills': found_functional[:10] if found_functional else ["None identified"],
        'Domain Skills': found_domain[:10] if found_domain else ["None identified"]
    }
    
    print(f"DEBUG: Skills extraction completed:")
    print(f"  Technical Skills: {len(found_technical)} found")
    print(f"  Functional Skills: {len(found_functional)} found")
    print(f"  Domain Skills: {len(found_domain)} found")
    
    return skills_table
def extract_experience(text):
    """Extract work experience"""
    experience_keywords = [
        'experience', 'work', 'employment', 'career', 'position',
        'job', 'role', 'worked', 'employed', 'served'
    ]
    
    lines = text.split('\n')
    experience_info = []
    
    # Look for years in format 2019-2022, 2019 - 2022, etc.
    year_pattern = r'\b(19|20)\d{2}\s*[-–]\s*(19|20)\d{2}|\b(19|20)\d{2}\s*[-–]\s*present\b'
    
    for i, line in enumerate(lines):
        line_lower = line.lower()
        
        # Check if line contains experience keywords or year patterns
        if (any(keyword in line_lower for keyword in experience_keywords) or 
            re.search(year_pattern, line, re.IGNORECASE)):
            
            exp_info = line.strip()
            
            # Try to get additional context from surrounding lines
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and len(next_line) < 100:
                    exp_info += f" - {next_line}"
            
            if exp_info and len(exp_info) > 15:  # Reasonable length
                experience_info.append(exp_info)
    
    return experience_info[:3] if experience_info else ["Not found"]

def load_certificate_keywords():
    """Load certificate keywords from external file"""
    keywords_file = os.path.join(os.path.dirname(__file__), 'keywords', 'cert_keywords.txt')
    certificate_keywords = []
    specific_certifications = []
    
    try:
        with open(keywords_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            # Skip empty lines and comments
            if not line or line.startswith('#'):
                continue
            
            # Add to appropriate list based on length and content
            # Shorter, generic terms go to certificate_keywords
            # Longer, specific certifications go to specific_certifications
            if len(line) <= 25 and any(word in line.lower() for word in ['certificate', 'certification', 'training', 'course', 'level']):
                certificate_keywords.append(line.lower())
            else:
                specific_certifications.append(line.lower())
                
    except FileNotFoundError:
        print("WARNING: cert_keywords.txt not found, using fallback keywords")
        # Fallback to basic keywords if file not found
        certificate_keywords = [
            'certificate', 'certification', 'certified', 'diploma', 'license', 'licensed',
            'accreditation', 'accredited', 'qualification', 'credential', 'award',
            'completion', 'training', 'course', 'workshop', 'seminar', 'bootcamp',
            'professional development', 'continuing education', 'level-', 'level -'
        ]
        specific_certifications = [
            'aws certified', 'microsoft certified', 'azure certified', 'pmp', 'scrum master'
        ]
    
    return certificate_keywords, specific_certifications

def extract_certificates(text):
    """Extract certificates and certifications from text"""
    print("DEBUG: Starting certificate extraction...")
    
    # Load certificate keywords from external file
    certificate_keywords, specific_certifications = load_certificate_keywords()
    
    print(f"DEBUG: Loaded {len(certificate_keywords)} general keywords and {len(specific_certifications)} specific certifications")
    
    # Section headers that indicate certificate sections
    section_headers = [
        'certificate', 'certification', 'training', 'course',  'seminar', 'certificates', 'trainings',
        'license', 'credential', 'achievement', 'award', 'honor'
    ]
    
    lines = text.split('\n')
    certificate_info = []
    in_certificate_section = False
    section_depth = 0
    
    print(f"DEBUG: Processing {len(lines)} lines for certificates...")
    
    for i, line in enumerate(lines):
        line_clean = line.strip()
        line_lower = line_clean.lower()
        
        # Skip empty lines
        if len(line_clean) < 2:
            continue
        
        print(f"DEBUG: Line {i+1}: '{line_clean}'")
        
        # Check if this line is a certificate section header
        is_section_header = False
        if len(line_clean) < 60:  # Section headers are usually short
            for header in section_headers:
                if (header in line_lower and 
                    len(line_clean.replace(header, '').strip()) < 10):  # Mostly just the header word
                    is_section_header = True
                    in_certificate_section = True
                    section_depth = 0
                    print(f"DEBUG: Found certificate section header at line {i+1}: '{line_clean}'")
                    break
        
        if is_section_header:
            continue
        
        # If we're in a certificate section, process more lines
        if in_certificate_section:
            section_depth += 1
            
            # Stop if we've gone too far or hit another major section
            if (section_depth > 20 or 
                any(other_section in line_lower for other_section in 
                    ['work experience', 'employment history', 'education', 'academic background', 
                     'technical skills', 'projects', 'summary', 'objective'])):
                in_certificate_section = False
                print(f"DEBUG: Exiting certificate section at line {i+1}")
                # Don't continue here - still check this line for certificates
            
            # Process potential certificate lines
            if in_certificate_section and len(line_clean) > 5:
                # Clean the line
                cert_line = line_clean
                
                # Remove bullet points and common prefixes
                for prefix in ['•', '▪', '○', '►', '→', '-', '*', ':', '1.', '2.', '3.', '4.', '5.']:
                    if cert_line.startswith(prefix):
                        cert_line = cert_line[len(prefix):].strip()
                
                # Skip if it's just a number or very generic
                if (len(cert_line) > 8 and 
                    not cert_line.lower() in ['certificates', 'certifications', 'training', 'courses']):
                    
                    formatted_cert = format_certificate_entry(cert_line)
                    if formatted_cert and formatted_cert not in certificate_info:
                        certificate_info.append(formatted_cert)
                        print(f"DEBUG: ADDED from section: {formatted_cert}")
        
        # Always check for specific certifications anywhere in the document
        matched_certifications = []
        for cert_pattern in specific_certifications:
            if cert_pattern in line_lower:
                matched_certifications.append(cert_pattern)
        
        if matched_certifications:
            print(f"DEBUG: Found specific certifications in line {i+1}: {matched_certifications}")
            cert_line = line_clean
            
            # Remove bullet points
            for prefix in ['•', '▪', '○', '►', '→', '-', '*', ':', '1.', '2.', '3.', '4.', '5.']:
                if cert_line.startswith(prefix):
                    cert_line = cert_line[len(prefix):].strip()
            
            formatted_cert = format_certificate_entry(cert_line)
            if formatted_cert and formatted_cert not in certificate_info:
                certificate_info.append(formatted_cert)
                print(f"DEBUG: ADDED specific certification: {formatted_cert}")
        
        # Check for general certificate patterns with strong indicators
        elif any(keyword in line_lower for keyword in certificate_keywords):
            # Look for strong certificate indicators
            strong_indicators = [
                'certificate of', 'certification in', 'certified in', 'diploma in',
                'license in', 'completion of', 'training in', 'course in', 'workshop on'
            ]
            
            if any(indicator in line_lower for indicator in strong_indicators):
                cert_line = line_clean
                
                # Remove bullet points
                for prefix in ['•', '▪', '○', '►', '→', '-', '*', ':', '1.', '2.', '3.', '4.', '5.']:
                    if cert_line.startswith(prefix):
                        cert_line = cert_line[len(prefix):].strip()
                
                if len(cert_line) > 10:
                    formatted_cert = format_certificate_entry(cert_line)
                    if formatted_cert and formatted_cert not in certificate_info:
                        certificate_info.append(formatted_cert)
                        print(f"DEBUG: ADDED general certificate: {formatted_cert}")
    
    print(f"DEBUG: Total certificates found: {len(certificate_info)}")
    
    # Remove duplicates and clean up
    unique_certificates = []
    seen = set()
    
    for cert in certificate_info:
        # Use the main certificate name for deduplication
        cert_name = cert.split(' - Year:')[0].strip().lower()
        if cert_name not in seen and len(cert_name) > 5:
            seen.add(cert_name)
            unique_certificates.append(cert)
    
    print(f"DEBUG: Final unique certificates: {unique_certificates}")
    return unique_certificates[:15] if unique_certificates else ["Not found"]

# Helper method for formatting certificate entries
def format_certificate_entry(cert_text):
    """Format a certificate entry with issuer extraction (year removed)"""
    if not cert_text or len(cert_text.strip()) < 5:
        return None
    
    cert_name = cert_text.strip()
    issuer = "Not specified"
    
    # Remove year from the certificate name if present
    cert_name = re.sub(r'\b(19|20)\d{2}\b', '', cert_name).strip()
    
    # Extract issuer from common patterns
    issuer_patterns = [
        r'\(([^)]+)\)',  # Text in parentheses
        r'from\s+([A-Za-z\s&,.-]+?)(?:\s|$)',  # "from XYZ"
        r'by\s+([A-Za-z\s&,.-]+?)(?:\s|$)',    # "by XYZ"
        r'-\s*([A-Za-z\s&,.-]+?)$',            # "Certificate - Issuer"
        r'issued by\s+([A-Za-z\s&,.-]+?)(?:\s|$)'  # "issued by XYZ"
    ]
    
    for pattern in issuer_patterns:
        issuer_match = re.search(pattern, cert_name, re.IGNORECASE)
        if issuer_match:
            potential_issuer = issuer_match.group(1).strip()
            if len(potential_issuer) > 2 and len(potential_issuer) < 50:
                issuer = potential_issuer
                # Remove the issuer part from the certificate name
                cert_name = re.sub(pattern, '', cert_name, flags=re.IGNORECASE).strip()
                break
    
    # Clean up certificate name
    cert_name = re.sub(r'\s+', ' ', cert_name).strip()
    cert_name = cert_name.rstrip('- ').strip()
    
    if len(cert_name) < 5:
        return None
    
    return f"{cert_name} - Issuer: {issuer}"

def parse_resume(file_path, file_extension):
    """Main function to parse resume and extract information"""
    
    # Extract text based on file type
    if file_extension == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif file_extension == 'docx':
        text = extract_text_from_docx(file_path)
    elif file_extension == 'txt':
        text = extract_text_from_txt(file_path)
    else:
        return None
    
    if not text.strip():
        return None
    
    # Extract information
    resume_data = {
        'name': extract_name(text),
        'email': extract_email(text),
        'phone': extract_phone(text),
        'education': extract_education(text),
        'skills': extract_skills(text),
        'experience': extract_experience(text),
        'certificates': extract_certificates(text)
    }
    
    return resume_data

@app.route('/')
def index():
    return render_template('resume_parser.html')

@app.route('/parse', methods=['POST'])
def parse_resume_endpoint():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type'}), 400
    
    tmp_file_path = None
    try:
        # Save file temporarily
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        # Create temporary file
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}')
        tmp_file_path = tmp_file.name
        tmp_file.close()  # Close the file handle
        
        # Save uploaded file to temporary location
        file.save(tmp_file_path)
        
        # Parse the resume
        resume_data = parse_resume(tmp_file_path, file_extension)
        
        if resume_data is None:
            return jsonify({'error': 'Could not extract text from file'}), 400
        
        return jsonify(resume_data)
    
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500
    
    finally:
        # Clean up temporary file
        if tmp_file_path and os.path.exists(tmp_file_path):
            try:
                os.unlink(tmp_file_path)
            except Exception:
                pass  # Ignore cleanup errors

if __name__ == '__main__':
    app.run(debug=True)
